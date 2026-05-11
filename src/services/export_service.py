import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.models.project import Project
from src.models.chapter import Chapter
from config.settings import get_settings


# Chinese font configuration for DOCX
FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_WESTERN = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADING1 = Pt(16)
FONT_SIZE_HEADING2 = Pt(14)
FONT_SIZE_TITLE = Pt(18)


def _set_run_font(run, font_name=FONT_BODY, size=FONT_SIZE_BODY, bold=False):
    """Set both Western and East-Asian fonts for a run."""
    run.font.name = FONT_WESTERN
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_heading(doc, text, level=1):
    """Add a heading with proper Chinese font."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        _set_run_font(run, FONT_HEADING, FONT_SIZE_TITLE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        _set_run_font(run, FONT_HEADING, FONT_SIZE_HEADING1, bold=True)
    elif level == 2:
        _set_run_font(run, FONT_HEADING, FONT_SIZE_HEADING2, bold=True)
    else:
        _set_run_font(run, FONT_HEADING, size=FONT_SIZE_BODY, bold=True)
    _set_paragraph_spacing(p, before=12, after=6)
    return p


def _add_body(doc, text):
    """Add body text with proper Chinese font."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, FONT_BODY, FONT_SIZE_BODY)
    _set_paragraph_spacing(p)
    # First line indent
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def _add_figure_placeholder(doc, figure):
    """Add a figure placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{figure['position_label']}] {figure['description']}")
    _set_run_font(run, FONT_BODY, Pt(10), bold=False)
    run.font.color.rgb = RGBColor(100, 100, 100)
    _set_paragraph_spacing(p, before=6, after=6)


def _markdown_to_docx_runs(paragraph, text):
    """Convert simple Markdown inline formatting to DOCX runs."""
    # Bold: **text**
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, FONT_BODY, FONT_SIZE_BODY, bold=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, FONT_BODY, FONT_SIZE_BODY)


def export_to_docx(project: Project) -> Path:
    """Export a project to a DOCX file with Chinese patent formatting."""
    settings = get_settings()
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Title
    _add_heading(doc, project.title or "专利交底书", level=0)
    _add_body(doc, f"专利类型：{project.patent_type}")
    _add_body(doc, f"技术领域：{project.domain}")
    _add_body(doc, f"撰写日期：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph()  # spacer

    # Chapters
    chapters = sorted(project.chapters, key=lambda c: c.chapter_order)
    for ch in chapters:
        if not ch.content:
            continue
        _add_heading(doc, f"{ch.chapter_order}. {ch.chapter_name}", level=1)

        # Parse content: split by lines, handle figure markers and markdown
        lines = ch.content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Check for figure markers
            if re.match(r"\[插入图\d+", line):
                for fig in (ch.figure_placeholders or []):
                    if fig["description"] in line or fig["position_label"] in line:
                        _add_figure_placeholder(doc, fig)
                        break
                i += 1
                continue

            # Markdown heading (###)
            if line.startswith("### "):
                _add_heading(doc, line[4:], level=3)
            elif line.startswith("## "):
                _add_heading(doc, line[3:], level=2)
            elif line.startswith("# "):
                _add_heading(doc, line[2:], level=1)
            elif re.match(r"^\d+[\.\)]\s", line):
                # Numbered list
                p = doc.add_paragraph()
                _markdown_to_docx_runs(p, line)
                _set_paragraph_spacing(p)
                p.paragraph_format.first_line_indent = Cm(0.74)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph()
                _markdown_to_docx_runs(p, "• " + line[2:])
                _set_paragraph_spacing(p)
                p.paragraph_format.left_indent = Cm(1.0)
            else:
                _add_body(doc, line)

            i += 1

        doc.add_paragraph()  # spacer between chapters

    # Save
    output_dir = Path(settings.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{project.title or '专利交底书'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    # Sanitize filename
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath


def export_to_markdown(project: Project) -> Path:
    """Export a project to a Markdown file."""
    settings = get_settings()
    output_dir = Path(settings.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    md = []
    md.append(f"# {project.title or '专利交底书'}")
    md.append(f"**专利类型**：{project.patent_type}  ")
    md.append(f"**技术领域**：{project.domain}  ")
    md.append(f"**撰写日期**：{datetime.now().strftime('%Y年%m月%d日')}  ")
    md.append("")
    md.append("---")
    md.append("")

    chapters = sorted(project.chapters, key=lambda c: c.chapter_order)
    for ch in chapters:
        if not ch.content:
            continue
        md.append(f"## {ch.chapter_order}. {ch.chapter_name}")
        md.append("")
        md.append(ch.content)
        md.append("")

        # Figure summary
        if ch.figure_placeholders:
            md.append("**本章附图**：")
            for fig in ch.figure_placeholders:
                md.append(f"- [{fig['position_label']}] {fig['description']}")
            md.append("")

    filename = f"{project.title or '专利交底书'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    filepath = output_dir / filename
    filepath.write_text("\n".join(md), encoding="utf-8")
    return filepath
